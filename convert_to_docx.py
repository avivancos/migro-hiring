#!/usr/bin/env python3
"""
Script para convertir el contrato de confidencialidad de Markdown a DOCX
"""
import re
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    print("Instalando python-docx...")
    import subprocess
    subprocess.check_call(["pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

def parse_markdown_to_docx(md_file_path, docx_file_path):
    """Convierte un archivo Markdown a DOCX con formato legal apropiado"""
    
    # Leer el archivo Markdown
    with open(md_file_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Crear documento Word
    doc = Document()
    
    # Configurar márgenes (márgenes estándar para documentos legales)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Dividir en líneas
    lines = content.split('\n')
    
    i = 0
    current_paragraph = None
    
    while i < len(lines):
        line = lines[i].strip()
        original_line = lines[i]  # Mantener espacios originales para indentación
        
        # Título principal (#)
        if line.startswith('# '):
            text = line[2:].strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            run.font.size = Pt(16)
            run.font.bold = True
            run.font.name = 'Times New Roman'
            doc.add_paragraph()  # Espacio después del título
            current_paragraph = None
        
        # Subtítulos (##)
        elif line.startswith('## '):
            text = line[3:].strip()
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.name = 'Times New Roman'
            doc.add_paragraph()
            current_paragraph = None
        
        # Encabezados de sección (REUNIDOS, EXPONEN, CLÁUSULAS)
        elif line in ['REUNIDOS', 'EXPONEN', 'CLÁUSULAS']:
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.name = 'Times New Roman'
            run.font.all_caps = True
            doc.add_paragraph()
            current_paragraph = None
        
        # Numeración de cláusulas (1., 2., etc.)
        elif re.match(r'^\d+\.\s+[A-ZÁÉÍÓÚÑ]', line):
            text = line
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.name = 'Times New Roman'
            current_paragraph = p
        
        # Subnumeración (1.1., 1.2., etc.)
        elif re.match(r'^\d+\.\d+\.', line):
            text = line
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            run = p.add_run(text)
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.name = 'Times New Roman'
            current_paragraph = p
        
        # Listas con letras (a), b), etc.)
        elif re.match(r'^\s*\([a-z]\)\s+', line):
            text = line.strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(-0.25)  # Sangría francesa
            run = p.add_run(text)
            run.font.size = Pt(11)
            run.font.name = 'Times New Roman'
            current_paragraph = None
        
        # Líneas de firma
        elif 'Por EL ASPIRANTE' in line or 'Por MIGRO' in line:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line)
            run.font.size = Pt(11)
            run.font.name = 'Times New Roman'
            current_paragraph = None
        
        # Líneas con guiones bajos (campos de firma)
        elif re.match(r'^[A-Za-z]+:\s*_{3,}', line):
            text = line
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(text)
            run.font.size = Pt(11)
            run.font.name = 'Times New Roman'
            current_paragraph = None
        
        # Línea vacía
        elif not line:
            doc.add_paragraph()
            current_paragraph = None
        
        # Texto normal - continuar párrafo anterior si existe
        else:
            if line:
                # Si hay un párrafo actual (cláusula numerada), continuar en él
                if current_paragraph is not None:
                    run = current_paragraph.add_run(' ' + line)
                    run.font.size = Pt(11)
                    run.font.name = 'Times New Roman'
                else:
                    # Crear nuevo párrafo
                    p = doc.add_paragraph()
                    run = p.add_run(line)
                    run.font.size = Pt(11)
                    run.font.name = 'Times New Roman'
                    current_paragraph = None
        
        i += 1
    
    # Guardar el documento
    doc.save(docx_file_path)
    print(f"Documento Word creado exitosamente: {docx_file_path}")

if __name__ == "__main__":
    # Rutas de archivos
    md_file = Path("src/legal/confidencialidad_agentes_vendedores.md")
    docx_file = Path("src/legal/confidencialidad_agentes_vendedores.docx")
    
    if not md_file.exists():
        print(f"Error: No se encontró el archivo {md_file}")
        exit(1)
    
    parse_markdown_to_docx(md_file, docx_file)

