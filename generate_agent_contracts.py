#!/usr/bin/env python3
"""
Script para generar contratos personalizados de colaboración de agentes
con formato verde Migro y exportación a PDF.
"""
import re
from pathlib import Path
from datetime import datetime
import subprocess
from typing import Optional

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ImportError:
    print("Instalando python-docx...")
    import subprocess
    subprocess.check_call(["pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

# Color verde Migro: RGB(22, 163, 74) = #16a34a
MIGRO_GREEN = RGBColor(22, 163, 74)
MIGRO_WHITE = RGBColor(255, 255, 255)

def normalize_spanish_date(raw_date: str) -> str:
    """Normaliza una fecha en español (ej: 28 de enero de 2026)."""
    return raw_date.strip().replace("  ", " ")

def convert_docx_to_pdf(docx_path: Path, pdf_path: Path) -> bool:
    """Convierte un DOCX a PDF usando docx2pdf o LibreOffice."""
    try:
        try:
            from docx2pdf import convert
        except ImportError:
            print("Instalando docx2pdf...")
            subprocess.check_call(["pip", "install", "docx2pdf"])
            from docx2pdf import convert
        convert(str(docx_path), str(pdf_path))
        print(f"PDF generado (docx2pdf): {pdf_path}")
        return True
    except Exception:
        pass

    try:
        if pdf_path.exists():
            # Evita que un PDF corrupto previo bloquee el fallback.
            try:
                pdf_path.unlink()
            except Exception:
                pass
        subprocess.check_call([
            "soffice",
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            str(pdf_path.parent),
            str(docx_path),
        ])
        # LibreOffice guarda con el mismo nombre base
        generated = pdf_path.parent / f"{docx_path.stem}.pdf"
        if generated.exists() and generated != pdf_path:
            generated.replace(pdf_path)
        print(f"PDF generado (LibreOffice): {pdf_path}")
        return True
    except Exception as exc:
        print(f"[ERROR] No se pudo convertir a PDF ({docx_path.name}): {exc}")
        return False

def add_signature_table(doc, agent_data, representante_migro):
    """Añade tabla de firmas con domicilios."""
    doc.add_paragraph("Firmado electrónicamente con HelloSign.")
    doc.add_paragraph()

    table = doc.add_table(rows=5, cols=2)
    table.autofit = False
    for row in table.rows:
        row.cells[0].width = Inches(3.2)
        row.cells[1].width = Inches(3.2)

    def set_cell(cell, text, bold=False):
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.name = "Times New Roman"
        run.font.bold = bold

    set_cell(table.cell(0, 0), "Por EL AGENTE", bold=True)
    set_cell(table.cell(0, 1), "Por MIGRO SERVICIOS Y REMESAS S.L.", bold=True)
    set_cell(table.cell(1, 0), f"Nombre: {agent_data['nombre']}")
    set_cell(table.cell(1, 1), f"Nombre: {representante_migro['nombre']}")
    set_cell(table.cell(2, 0), f"{agent_data.get('tipo_doc', 'DNI/NIE')}: {agent_data.get('documento', '')}")
    set_cell(table.cell(2, 1), f"NIF: {representante_migro['nif']}")
    set_cell(table.cell(3, 0), f"Domicilio: {agent_data.get('domicilio', '')}")
    set_cell(table.cell(3, 1), "Domicilio: C/ Libreros, nº 54, 37008, Salamanca")
    set_cell(table.cell(4, 0), "Firma: ____________________________")
    set_cell(table.cell(4, 1), "Firma: ____________________________")

def parse_markdown_to_docx(md_file_path, docx_file_path, agent_data, fecha_firma: Optional[str] = None):
    """Convierte el contrato Markdown a DOCX con datos del agente y formato verde"""
    
    # Leer el archivo Markdown
    with open(md_file_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Datos del representante de Migro
    representante_migro = {
        'nombre': 'Agustin Vivancos Cabrillana',
        'nif': '25597879Y',
        'cargo': 'Representante Legal'
    }
    
    # Reemplazar datos del agente
    content = content.replace('[NOMBRE DEL ASPIRANTE]', agent_data['nombre'])
    content = content.replace('[AGENTE DE VENTAS]', agent_data['nombre'])
    
    # Reemplazar documento del agente - buscar el patrón específico
    content = re.sub(
        r'con DNI/NIE \[●\]',
        f"con {agent_data.get('tipo_doc', 'DNI/NIE')} {agent_data.get('documento', '[●]')}",
        content,
    )
    
    # Reemplazar domicilio del agente
    content = re.sub(r'domicilio en \[●\]', f"domicilio en {agent_data.get('domicilio', '[●]')}", content)
    
    # Reemplazar datos del representante de Migro - reemplazar toda la línea completa
    migro_line_pattern = r'Y de otra, MIGRO SERVICIOS Y REMESAS S\.L\., en adelante "MIGRO", con CIF B22759765, domicilio en C/ Libreros, nº 54, 37008, Salamanca, actuando en su nombre y representación \[Nombre y Apellidos del representante\], con NIF \[●\], en su calidad de \[cargo\]\.'
    migro_line_replacement = f'Y de otra, MIGRO SERVICIOS Y REMESAS S.L., en adelante "MIGRO", con CIF B22759765, domicilio en C/ Libreros, nº 54, 37008, Salamanca, actuando en su nombre y representación {representante_migro["nombre"]}, con NIF {representante_migro["nif"]}, en su calidad de {representante_migro["cargo"]}.'
    content = re.sub(migro_line_pattern, migro_line_replacement, content)
    
    # Reemplazar zona de firmas completa con marcador para tabla
    firma_section = "[FIRMAS_TABLA]"
    
    # Buscar y reemplazar desde "Por EL ASPIRANTE" hasta el final del documento
    # Dividir el contenido en dos partes: antes y después de las firmas
    lines = content.split('\n')
    before_firma = []
    found_firma_start = False
    
    for i, line in enumerate(lines):
        if 'Por EL ASPIRANTE' in line or 'Por EL AGENTE' in line:
            found_firma_start = True
            break
        before_firma.append(line)
    
    if found_firma_start:
        # Reconstruir el contenido con la nueva sección de firmas
        content = '\n'.join(before_firma) + '\n\n' + firma_section
    else:
        # Si no se encontró, intentar reemplazo con regex
        firma_pattern = r'(Por EL (ASPIRANTE|AGENTE).*?Firma: _{10,}.*?Firma: _{10,})'
        content = re.sub(firma_pattern, firma_section, content, flags=re.DOTALL)
    
    # Fecha actual o fecha de firma forzada
    if fecha_firma:
        fecha_actual = normalize_spanish_date(fecha_firma)
    else:
        fecha_actual = datetime.now().strftime("%d de %B de %Y")
        meses_es = {
            'January': 'enero', 'February': 'febrero', 'March': 'marzo',
            'April': 'abril', 'May': 'mayo', 'June': 'junio',
            'July': 'julio', 'August': 'agosto', 'September': 'septiembre',
            'October': 'octubre', 'November': 'noviembre', 'December': 'diciembre'
        }
        for eng, esp in meses_es.items():
            fecha_actual = fecha_actual.replace(eng, esp)
    
    content = content.replace('[Ciudad], a [fecha].', f"Salamanca, a {fecha_actual}.")
    
    # Crear documento Word
    doc = Document()
    
    # Configurar márgenes
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Header verde Migro usando tabla sin bordes
    header_table = doc.add_table(rows=2, cols=1)
    header_table.style = 'Light Grid Accent 1'
    
    # Ocultar bordes de la tabla
    tbl = header_table._tbl
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'nil')
        border.set(qn('w:sz'), '0')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        tblBorders.append(border)
    tbl.tblPr.append(tblBorders)
    
    # Primera fila: Título principal
    cell1 = header_table.rows[0].cells[0]
    cell1.vertical_alignment = 1  # Center
    para1 = cell1.paragraphs[0]
    para1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para1.paragraph_format.space_after = Pt(0)
    run1 = para1.add_run('MIGRO SERVICIOS Y REMESAS S.L.')
    run1.font.size = Pt(18)
    run1.font.bold = True
    run1.font.color.rgb = MIGRO_WHITE
    run1.font.name = 'Arial'
    
    # Fondo verde para la primera celda
    tcPr = cell1._element.get_or_add_tcPr()
    shd1 = OxmlElement('w:shd')
    shd1.set(qn('w:fill'), '16a34a')
    tcPr.append(shd1)
    
    # Segunda fila: Subtítulo
    cell2 = header_table.rows[1].cells[0]
    cell2.vertical_alignment = 1  # Center
    para2 = cell2.paragraphs[0]
    para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para2.paragraph_format.space_after = Pt(0)
    run2 = para2.add_run('CIF: B22759765 | C/ Libreros, nº 54, 37008 - Salamanca')
    run2.font.size = Pt(10)
    run2.font.color.rgb = MIGRO_WHITE
    run2.font.name = 'Arial'
    
    # Fondo verde para la segunda celda
    tcPr2 = cell2._element.get_or_add_tcPr()
    shd2 = OxmlElement('w:shd')
    shd2.set(qn('w:fill'), '16a34a')
    tcPr2.append(shd2)
    
    doc.add_paragraph()  # Espacio después del header
    
    # Dividir en líneas
    lines = content.split('\n')
    
    i = 0
    current_paragraph = None
    
    while i < len(lines):
        line = lines[i].strip()
        
        # Título principal (#)
        if line.startswith('# '):
            text = line[2:].strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            run.font.size = Pt(16)
            run.font.bold = True
            run.font.name = 'Times New Roman'
            doc.add_paragraph()
            current_paragraph = None
        
        # Subtítulos (##)
        elif line.startswith('## '):
            text = line[3:].strip()
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.name = 'Times New Roman'
            doc.add_paragraph()
            current_paragraph = None
        
        # Encabezados de sección (REUNIDOS, EXPONEN, CLÁUSULAS)
        elif line in ['REUNIDOS', 'EXPONEN', 'CLÁUSULAS']:
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.name = 'Times New Roman'
            run.font.all_caps = True
            doc.add_paragraph()
            current_paragraph = None
        
        # Numeración de cláusulas (1., 2., etc.)
        elif re.match(r'^\d+\.\s+[A-ZÁÉÍÓÚÑ]', line):
            text = line
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.name = 'Times New Roman'
            current_paragraph = p
        
        # Subnumeración (1.1., 1.2., etc.)
        elif re.match(r'^\d+\.\d+\.', line):
            text = line
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            run = p.add_run(text)
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.name = 'Times New Roman'
            current_paragraph = p
        
        # Listas con letras (a), b), etc.)
        elif re.match(r'^\s*\([a-z]\)\s+', line):
            text = line.strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(-0.25)
            run = p.add_run(text)
            run.font.size = Pt(11)
            run.font.name = 'Times New Roman'
            current_paragraph = None
        
        # Marcador de tabla de firmas
        elif line == '[FIRMAS_TABLA]':
            add_signature_table(doc, agent_data, representante_migro)
            current_paragraph = None
        
        # Líneas con guiones bajos (campos de firma)
        elif re.match(r'^[A-Za-z]+:\s*_{3,}', line):
            text = line
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(text)
            run.font.size = Pt(11)
            run.font.name = 'Times New Roman'
            current_paragraph = None
        
        # Línea vacía
        elif not line:
            doc.add_paragraph()
            current_paragraph = None
        
        # Texto normal
        else:
            if line:
                if current_paragraph is not None:
                    run = current_paragraph.add_run(' ' + line)
                    run.font.size = Pt(11)
                    run.font.name = 'Times New Roman'
                else:
                    p = doc.add_paragraph()
                    run = p.add_run(line)
                    run.font.size = Pt(11)
                    run.font.name = 'Times New Roman'
                    current_paragraph = None
        
        i += 1
    
    # Guardar el documento
    doc.save(docx_file_path)
    print(f"Contrato generado: {docx_file_path}")

# Datos de los agentes
agentes = [
    {
        'nombre': 'Erick Ezequiel Alvarez Legorreta',
        'documento': 'AALE950416HMCLGR00',
        'tipo_doc': 'CURP',
        'domicilio': 'Av de los Maestros 314, Colonia Doctores, C.P. 50060, Toluca, México',
        'archivo': 'erick_ezequiel_alvarez_legorreta'
    },
    {
        'nombre': 'Gabriela Calderon Alvarado',
        'documento': 'CAAG861227MMNLLB07',
        'tipo_doc': 'CURP',
        'domicilio': 'Blvd. Sauces #161, Fraccionamiento Rinconada los Sauces, Tarimbaro, Michoacán, México',
        'archivo': 'gabriela_calderon_alvarado'
    },
    {
        'nombre': 'Brenda Montserrat Fanny Moreno Tovar',
        'documento': 'IDMEX1832667425',
        'tipo_doc': 'INE',
        'domicilio': 'Calzada de las águilas #934, Col. Ampliación las águilas, C.P. 01759, Alcaldía Álvaro Obregón, México',
        'archivo': 'brenda_montserrat_fanny_moreno_tovar'
    },
    {
        'nombre': 'Sonia Alejandra Cisnero',
        'documento': '26245201',
        'tipo_doc': 'DNI',
        'domicilio': 'Sebastián El Cano 147, San Rafael, Mendoza, Argentina',
        'archivo': 'sonia_alejandra_cisnero'
    }
]

if __name__ == "__main__":
    md_file = Path("src/legal/agente_ventas_agreement.md")
    
    if not md_file.exists():
        print(f"Error: No se encontró el archivo {md_file}")
        exit(1)
    
    # Crear directorio para los contratos si no existe
    output_dir = Path("src/legal/contratos_agentes")
    output_dir.mkdir(exist_ok=True)
    
    print("Generando contratos personalizados...\n")
    
    exitosos = 0
    fallidos = []
    
    fecha_firma = "28 de enero de 2026"
    
    for agente in agentes:
        docx_file = output_dir / f"convenio_colaboracion_{agente['archivo']}.docx"
        pdf_file = output_dir / f"convenio_colaboracion_{agente['archivo']}.pdf"
        try:
            parse_markdown_to_docx(md_file, docx_file, agente, fecha_firma=fecha_firma)
            convert_docx_to_pdf(docx_file, pdf_file)
            exitosos += 1
        except PermissionError as e:
            print(f"  [AVISO] Archivo bloqueado: {agente['nombre']}")
            fallidos.append(agente['nombre'])
        except Exception as e:
            print(f"  [ERROR] Error generando contrato para {agente['nombre']}: {e}")
            fallidos.append(agente['nombre'])
    
    print(f"\n{'='*60}")
    print(f"Contratos generados exitosamente: {exitosos}/{len(agentes)}")
    if fallidos:
        print(f"Contratos no generados: {len(fallidos)}")
        for nombre in fallidos:
            print(f"  - {nombre}")
    print(f"\nUbicación: {output_dir}")

